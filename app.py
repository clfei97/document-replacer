import tkinter as tk
import sys, os
import csv
import xlrd
from docx import Document
from datetime import datetime
import math

window = tk.Tk()
window.title("Doc replacer custom ver 1.0 ")

if getattr(sys, 'frozen', False):
    accExcPath = os.path.join(os.path.dirname(sys.executable),"csv","accList.xlsx")
    cardExcPath = os.path.join(os.path.dirname(sys.executable),"csv","cardList.xlsx")
    accTemplatePath = os.path.join(os.path.dirname(sys.executable),"template","jdb_account.docx")
    cardTemplatePath = os.path.join(os.path.dirname(sys.executable),"template","jdb_card.docx")
    saveAccPath = os.path.join(os.path.dirname(sys.executable),"generate acc")
    saveCardPath = os.path.join(os.path.dirname(sys.executable),"generate card")
else:
    accExcPath = os.path.join(sys.path[0],"csv","accList.xlsx")
    cardExcPath = os.path.join(sys.path[0],"csv","cardList.xlsx")
    accTemplatePath = os.path.join(sys.path[0],"template","jdb_account.docx")
    cardTemplatePath = os.path.join(sys.path[0],"template","jdb_card.docx")
    saveAccPath = os.path.join(sys.path[0],"generate acc")
    saveCardPath = os.path.join(sys.path[0],"generate card")

def startAccProcess():
    text.configure(text="Please wait......", fg="red")
    #open file here
    try:
        wb = xlrd.open_workbook(accExcPath)
        sheet = wb.sheet_by_index(0)
        sheet.cell_value(0, 0)
        replaceText = []
        line_count = 0
        for x in range(sheet.nrows):
            doc = Document(accTemplatePath)
            if x == 0:
                for y in range(sheet.ncols):
                    replaceText.append(sheet.cell_value(x,y))
            else:
                for y in range(sheet.ncols):
                    if y == 0:
                        filename = sheet.cell_value(x,y)
                    else:
                        for p in doc.paragraphs:
                            if p.text.find(replaceText[y])>=0:
                                if sheet.cell_value(x,y) is not None:
                                    if y == 2:
                                        try:
                                            python_date = datetime(*xlrd.xldate_as_tuple(sheet.cell_value(x,y), 0))
                                            p.text=p.text.replace(replaceText[y],str(python_date.strftime("%d/%m/%Y")))
                                        except:                                           
                                            p.text=p.text.replace(replaceText[y],str(sheet.cell_value(x,y)))
                                    else:
                                        value = sheet.cell_value(x,y)
                                        if(type(sheet.cell_value(x,y))==float):
                                            value = math.trunc(sheet.cell_value(x,y))
                                        p.text=p.text.replace(replaceText[y],str(value))
                    newSavePath = os.path.join(saveAccPath, f"{filename}.docx")
                    doc.save(newSavePath)
                line_count += 1
        text.configure(text=f"Processed {line_count} documents for account", fg="green")
    except Exception as e:
        text.configure(text=e, fg="red")



def startCardProcess():
    text.configure(text="Please wait......", fg="red")
    #open file here
    try:
        wb = xlrd.open_workbook(cardExcPath)
        sheet = wb.sheet_by_index(0)
        sheet.cell_value(0, 0)
        replaceText = []
        line_count = 0
        for x in range(sheet.nrows):
            doc = Document(cardTemplatePath)
            if x == 0:
                for y in range(sheet.ncols):
                    replaceText.append(sheet.cell_value(x,y))
            else:
                for y in range(sheet.ncols):
                    if y == 0:
                            filename = sheet.cell_value(x,y)
                    else:
                        for p in doc.paragraphs:
                            if p.text.find(replaceText[y])>=0:
                                if sheet.cell_value(x,y) is not None:
                                    if y == 3 or y == 5 or y == 7:
                                        try:
                                            python_date = datetime(*xlrd.xldate_as_tuple(sheet.cell_value(x,y), 0))
                                            p.text=p.text.replace(replaceText[y],str(python_date.strftime("%d/%m/%Y")))
                                        except:
                                            p.text=p.text.replace(replaceText[y],str(sheet.cell_value(x,y)))
                                    else:
                                        value = sheet.cell_value(x,y)
                                        if(type(sheet.cell_value(x,y))==float):
                                            value = math.trunc(sheet.cell_value(x,y))
                                        p.text=p.text.replace(replaceText[y],str(value))
                    newCardSavePath = os.path.join(saveCardPath, f"{filename}.docx")
                    doc.save(newCardSavePath)
                line_count += 1
        text.configure(text=f"Processed {line_count} documents for account", fg="green")
    except Exception as e:
        text.configure(text=e, fg="red")
            

"""
    with open(csvPath, "r", encoding="gbk") as csv_file:
        csv_reader = csv.reader(csv_file, delimiter=',')
        line_count = 0
        for row in csv_reader:
            if line_count == 0:
                replaceText = row
                line_count += 1
            else:
                for i in range(len(replaceText)):
                    if i == 0:
                        filename = row[i].decode("utf-8")
                    else:
                        for p in doc.paragraphs:
                            if p.text.find(replaceText[i])>=0:
                                p.text=p.text.replace(replaceText[i],row[i])
                    newSavePath = os.path.join(savePath, f"{filename}.docx")
                    doc.save(newSavePath)
                line_count += 1
        text.configure(text=f"Processed {line_count} documents", fg="green")
        """

canvas = tk.Canvas(window, height=700, width=700)
canvas.pack()
text = tk.Label(text="welcome", font=("Courier", 14))
text.place(relx = 0.5, rely = 0.8, anchor = 'center')
startAccBtn = tk.Button(window, text="Start for account", font=("Courier", 30), fg="white", bg="blue", command=startAccProcess)
startAccBtn.place(relx = 0.5, rely = 0.3, anchor = 'center')
startAccBtn = tk.Button(window, text="Start for card", font=("Courier", 30), fg="white", bg="blue", command=startCardProcess)
startAccBtn.place(relx = 0.5, rely = 0.5, anchor = 'center')

window.mainloop()